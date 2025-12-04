#!/usr/bin/env python3
"""
Convert TECHNICAL_JOURNEY.md to Word document (.docx)
Requires: pip install python-docx markdown
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def markdown_to_docx(md_file, docx_file):
    """Convert Markdown file to Word document"""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 80)
        
        # Tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Check if next line is table separator
            if '---' in lines[i + 1] or '|' in lines[i + 1]:
                table_data = []
                # Collect table rows
                j = i
                while j < len(lines) and '|' in lines[j]:
                    if '---' not in lines[j]:
                        row = [cell.strip() for cell in lines[j].split('|')[1:-1]]
                        table_data.append(row)
                    j += 1
                
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_data in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_data
                            # Make header row bold
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    
                    i = j - 1
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph(''.join(code_lines))
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
        
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Handle bold text
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_paragraph(text, style='List Number')
        
        # Regular paragraphs
        else:
            # Clean up markdown formatting
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            
            # Handle links [text](url)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            
            p = doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"✅ Converted {md_file} to {docx_file}")

if __name__ == '__main__':
    md_file = Path(__file__).parent / 'TECHNICAL_JOURNEY.md'
    docx_file = Path(__file__).parent / 'TECHNICAL_JOURNEY.docx'
    
    if not md_file.exists():
        print(f"❌ Error: {md_file} not found")
        exit(1)
    
    try:
        markdown_to_docx(md_file, docx_file)
        print(f"\n📄 Word document created: {docx_file}")
        print("   You can now open it in Microsoft Word and apply final formatting.")
    except ImportError:
        print("❌ Error: Required packages not installed")
        print("\nInstall with:")
        print("  pip install python-docx")
        exit(1)
    except Exception as e:
        print(f"❌ Error: {e}")
        exit(1)

